from datetime import time
import urllib.request
import win32com.client
import os
import time



from docx import Document
import pandas as pd

class BREMERHAVEN_CMA3PF():
    


    def read_docx_table(document,table_num=1,nheader=1):
        table = document.tables[table_num-1]
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data)
        if nheader == 1:
            df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)
        elif nheader ==2:
            outside_col, inside_col = df.iloc[0],df.ilocp[1]
            hier_index = pd.MultiIndex.from_tuples(list(zip(outside_col,inside_col)))
            df = pd.DataFrame(data,columns=hier_index).drop(df.index[[0,1]]).reset_index(drop=True)
        elif nheader > 2:
            print("More that two header currenty not supported")
            df= pd.DataFrame()
        return df

    pdf_path = ""
    def download_file(download_url, filename):
        response = urllib.request.urlopen(download_url)    
        file = open(filename + ".pdf", 'wb')
        # file = open(filename + ".pdf", 'w+')
        file.write(response.read())
        file.close()
    

    def BREMERHAVEN_t3PF():

        BREMERHAVEN_CMA3PF.download_file("https://tracking.ntb.eu/NTBSailingList.pdf", "NTBSailingList")
        
        word = win32com.client.Dispatch("Word.Application")
        word.visible = 0

        doc_pdf = "NTBSailingList.pdf"
        input_file = os.path.abspath(doc_pdf)

        wb = word.Documents.Open(input_file)
        
        output_file = os.path.abspath(doc_pdf[0:-4] + ".docx".format())
        # output_file = "NTBSailingList.docx"
        # wb.SaveAs(output_file,16)
        wb.SaveAs(output_file)
        wb.Close()
        word.Quit()
        # wb.Close()

        time.sleep(5)

        nDf = pd.DataFrame()
        document = Document("NTBSailingList.docx")
        table_num=1
        table_count = len(document.tables)
        for i in range(table_count):
            table_num =i
            nheader=1
            df = BREMERHAVEN_CMA3PF.read_docx_table(document,table_num,nheader)
            nDf = nDf.append(df)    

        # try:
        #     document = docx.Document("NTBSailingList.docx")
        #     # document = open('NTBSailingList.docx')
        #     # document = Document(output_file)
            
        # except:
        #     document = docx.Document("NTBSailingList.docx")
        #     # document = open('NTBSailingList.docx')
        #     # document.save("NTBSailingList.docx")

        # tables = document.tables
        # df = pd.DataFrame()

        # for table in document.tables:
        #     for row in table.rows:
        #         text = [cell.text for cell in row.cells]
        #         df = df.append([text], ignore_index=True)
        file_name =  pd.ExcelWriter('Terminals_Data.xlsx', engine='openpyxl',mode='a',if_sheet_exists='replace')
        nDf.to_excel(file_name,index=False,sheet_name='Germany_CMA3PF')
        file_name.save()     
                     
        os.remove("NTBSailingList.docx")
        os.remove("NTBSailingList.pdf")
        return        